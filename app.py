import streamlit as st
import faiss
import os
from io import BytesIO
from docx import Document
import numpy as np
from langchain_community.document_loaders import WebBaseLoader
from PyPDF2 import PdfReader
from langchain.chains import RetrievalQA
from langchain.text_splitter import CharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.docstore.in_memory import InMemoryDocstore
from langchain_huggingface import HuggingFaceEndpoint
from secret_api_keys import huggingface_api_key 

# Set HuggingFace API token in the environment
os.environ['HUGGINGFACEHUB_API_TOKEN'] = huggingface_api_key


def process_input(input_type, input_data):
    """
    Processes various input types (Link, PDF, Text, DOCX, TXT) and creates a FAISS vectorstore.

    Args:
        input_type (str): The type of input (e.g., "PDF", "Text").
        input_data: The input data, either a file-like object or a string.

    Returns:
        FAISS vectorstore containing embedded documents.
    """
    loader = None
    if input_type == "Link":
        loader = WebBaseLoader(input_data)  # Load web content
        documents = loader.load()  # Fetch content from the link(s)
    elif input_type == "PDF":
        # Process PDF file input
        if isinstance(input_data, BytesIO):
            pdf_reader = PdfReader(input_data)
        elif isinstance(input_data, UploadedFile):
            pdf_reader = PdfReader(BytesIO(input_data.read()))
        else:
            raise ValueError("Invalid input data for PDF")
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        documents = text
    elif input_type == "Text":
        # Directly use input text
        if isinstance(input_data, str):
            documents = input_data
        else:
            raise ValueError("Expected a string for 'Text' input type.")
    elif input_type == "DOCX":
        # Process DOCX file input
        if isinstance(input_data, BytesIO):
            doc = Document(input_data)
        elif isinstance(input_data, UploadedFile):
            doc = Document(BytesIO(input_data.read()))
        else:
            raise ValueError("Invalid input data for DOCX")
        text = "\n".join([para.text for para in doc.paragraphs])
        documents = text
    elif input_type == "TXT":
        # Process plain text file input
        if isinstance(input_data, BytesIO):
            text = input_data.read().decode('utf-8')
        elif isinstance(input_data, UploadedFile):
            text = str(input_data.read().decode('utf-8'))
        else:
            raise ValueError("Invalid input data for TXT")
        documents = text
    else:
        raise ValueError("Unsupported input type")

    # Split text into smaller chunks for embedding
    text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    if input_type == "Link":
        texts = text_splitter.split_documents(documents)
        texts = [str(doc.page_content) for doc in texts]  # Extract text content from Document objects
    else:
        texts = text_splitter.split_text(documents)

    # Initialize HuggingFace embedding model
    model_name = "sentence-transformers/all-mpnet-base-v2"
    model_kwargs = {'device': 'cpu'} 
    encode_kwargs = {'normalize_embeddings': False}
    hf_embeddings = HuggingFaceEmbeddings(
        model_name=model_name,
        model_kwargs=model_kwargs,
        encode_kwargs=encode_kwargs
    )

    # Initialize FAISS index
    sample_embedding = np.array(hf_embeddings.embed_query("sample text"))  # Sample to determine dimensions
    dimension = sample_embedding.shape[0]
    index = faiss.IndexFlatL2(dimension)  # Use L2 distance metric

    # Create FAISS vectorstore
    vector_store = FAISS(
        embedding_function=hf_embeddings.embed_query,
        index=index,
        docstore=InMemoryDocstore(),  # In-memory storage for small-scale use
        index_to_docstore_id={},
    )
    vector_store.add_texts(texts)  # Add processed text chunks to the vectorstore
    return vector_store


def answer_question(vectorstore, query):
    """
    Uses a pre-trained HuggingFace model to answer a user query based on the vectorstore.

    Args:
        vectorstore: FAISS vectorstore containing embedded documents.
        query (str): The user query.

    Returns:
        str: The answer to the query.
    """
    llm = HuggingFaceEndpoint(
        repo_id='meta-llama/Meta-Llama-3-8B-Instruct',  # Model repo
        token=huggingface_api_key,  # API token
        temperature=0.6  # Control output creativity
    )
    qa = RetrievalQA.from_chain_type(llm=llm, retriever=vectorstore.as_retriever())

    # Invoke the QA model with the user query
    answer = qa.invoke({"query": query})
    return answer


def main():
    """
    Main function to run the Streamlit app for RAG (Retrieval-Augmented Generation) Q&A.
    """
    st.title('RAG Q&A APP')

    # Input type selection
    input_type = st.selectbox("Input Type", ["Link", "PDF", "Text", "DOCX", "TXT"])
    input_data = None

    if input_type == "Link":
        number_input = st.number_input(min_value=1, max_value=20, step=1, label="Enter the number of Links")
        input_data = []
        for i in range(number_input):
            url = st.sidebar.text_input(f"URL {i+1}")
            input_data.append(url)
    elif input_type == "Text":
        input_data = st.text_input("Enter the text")
    elif input_type == 'PDF':
        input_data = st.file_uploader("Upload a PDF file", type=["pdf"])
    elif input_type == 'TXT':
        input_data = st.file_uploader("Upload a text file", type=['txt'])
    elif input_type == 'DOCX':
        input_data = st.file_uploader("Upload a DOCX file", type=['docx', 'doc'])

    # Process input and generate vectorstore
    if st.button("Proceed"):
        try:
            vectorstore = process_input(input_type, input_data)
            st.session_state["vectorstore"] = vectorstore  # Store in session state
        except Exception as e:
            st.error(f"Error processing input: {e}")

    # Answering user queries
    if "vectorstore" in st.session_state:
        query = st.text_input("Ask your question")
        if st.button("Submit"):
            try:
                answer = answer_question(st.session_state["vectorstore"], query)
                st.write(answer)
            except Exception as e:
                st.error(f"Error answering question: {e}")


if __name__ == '__main__':
    main()
